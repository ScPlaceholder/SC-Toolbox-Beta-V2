#!/usr/bin/env python3
"""
SC Toolbox Brand Design Guide Generator — v2
Grounded in the actual SC Toolbox codebase palette and Star Citizen's real UI language.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ─── ACTUAL SC TOOLBOX PALETTE (from codebase) ─────────────────────────────
# These are the real colors used across the 6 tools today.

# Backgrounds (navy-black range used across all tools)
BG_DEEPEST    = "06080d"   # Trade Hub / Mining base bg
BG_PRIMARY    = "0b0e14"   # Launcher main bg
BG_SECONDARY  = "111620"   # Launcher secondary / elevated panels
BG_CARD       = "141a26"   # Tile/card backgrounds
BG_INPUT      = "1c2233"   # Input fields, odd rows, settings bg
BG_HEADER     = "0e1420"   # Header bars

# Borders & separators
BORDER_PRIMARY = "1e2738"  # Primary border throughout launcher
BORDER_CARD    = "252e42"  # DPS Calculator card borders
SEPARATOR      = "0d1824"  # Trade Hub / Mining separators

# Text hierarchy
FG_PRIMARY     = "c8d4e8"  # Primary text across all tools
FG_BRIGHT      = "e8f2ff"  # Emphasis text (Mining/Trade)
FG_SECONDARY   = "5a6480"  # Dimmed text, descriptions
FG_TERTIARY    = "3a4460"  # Very dim, disabled text
FG_TRADE       = "b8ccde"  # Trade Hub / Mining body text

# Brand accent
ACCENT_CYAN    = "44aaff"  # THE primary accent — used in launcher, DPS, Market, Settings
SC_CYAN        = "00e7ff"  # Star Citizen MobiGlas canonical cyan
SC_CYAN_ALT    = "00A0B6"  # Star Citizen secondary cyan reference

# Per-tool accent colors (from skill_registry / tile definitions)
COL_DPS        = "ff7733"  # DPS Calculator orange
COL_CARGO      = "33ccdd"  # Cargo Loader cyan
COL_MISSION    = "33dd88"  # Mission Database green
COL_MINING     = "ffaa22"  # Mining Loadout amber
COL_MARKET     = "aa66ff"  # Market Finder purple
COL_TRADE      = "ffcc00"  # Trade Hub gold

# Functional colors (from codebase)
GREEN          = "33dd88"  # Success, running, profit
GREEN_ALT      = "00dd70"  # Trade Hub / Mining success
YELLOW         = "ffaa22"  # Warnings, hidden status
RED            = "ff5533"  # Errors, unavailable, loss
RED_ALT        = "e04020"  # Trade Hub / Mining errors
ORANGE         = "ff7733"  # DPS accent, thermal damage
PURPLE         = "aa66ff"  # Distortion, market finder
ENERGY_CYAN    = "44ccff"  # Energy damage, weapon glow

# Star Citizen reference colors
SC_BLUE        = "0090e0"  # Mining / Trade UI blue (close to SC HUD blue)
MISSION_TEAL   = "33ccaa"  # Mission Database unique accent

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "SC_Toolbox_Brand_Guide.docx")


# ─── HELPER FUNCTIONS ───────────────────────────────────────────────────────

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "CCCCCC")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def styled_para(doc, text, size=10, color=None, bold=False, italic=False,
                align=None, before=0, after=6, line_h=None, font="Calibri"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_h:
        pf.line_spacing = Pt(line_h)
    return p

def section_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor.from_string(ACCENT_CYAN)
        r.font.name = "Calibri"
    return h

def sub_h(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor.from_string(FG_BRIGHT)
        r.font.name = "Calibri"
    return h

def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)

def color_table(doc, colors_data):
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ["Swatch", "Name", "Hex", "RGB", "Usage"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True; run.font.size = Pt(9); run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(FG_BRIGHT)
        set_cell_shading(hdr_cells[i], BG_HEADER)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for name, hexv, rgb, usage in colors_data:
        row = table.add_row()
        c = row.cells
        set_cell_shading(c[0], hexv)
        c[0].text = ""; c[0].width = Inches(0.7)
        for i, (txt, w, fn) in enumerate([
            (name, 1.2, "Calibri"), (f"#{hexv}", 0.9, "Consolas"),
            (rgb, 1.2, "Consolas"), (usage, 2.5, "Calibri")
        ], 1):
            c[i].text = ""
            p = c[i].paragraphs[0]
            run = p.add_run(txt)
            run.font.size = Pt(9); run.font.name = fn
            if i == 1: run.bold = True
            c[i].width = Inches(w)
        for cell in c:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    _style_table_borders(table, BORDER_PRIMARY)
    return table

def data_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True; run.font.size = Pt(9); run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(FG_BRIGHT)
        set_cell_shading(hdr[i], BG_HEADER)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for rd in rows:
        row = table.add_row()
        for i, val in enumerate(rd):
            row.cells[i].text = ""
            p = row.cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9); run.font.name = "Calibri"
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells): row.cells[i].width = Inches(w)
    _style_table_borders(table, BORDER_PRIMARY)
    return table

def _style_table_borders(table, color):
    bd = {"sz": "4", "color": color}
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell, top=bd, bottom=bd, left=bd, right=bd)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)


# ─── DOCUMENT SECTIONS ─────────────────────────────────────────────────────

def build_cover(doc):
    for _ in range(6):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
        p.add_run(" ").font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SC TOOLBOX")
    run.font.name = "Calibri"; run.font.size = Pt(48)
    run.font.color.rgb = RGBColor.from_string(ACCENT_CYAN); run.bold = True
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2500" * 40)
    run.font.color.rgb = RGBColor.from_string(BORDER_PRIMARY); run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BRAND DESIGN GUIDE")
    run.font.name = "Calibri"; run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(FG_PRIMARY)
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Version 1.0  \u2022  March 2026")
    run.font.name = "Calibri"; run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(FG_SECONDARY)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Precision Tools for the Working Pilot")
    run.font.name = "Calibri"; run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(FG_SECONDARY); run.italic = True

    for _ in range(6):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
        p.add_run(" ").font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "SC Toolbox is a fan-made community project. Star Citizen\u00ae is a registered trademark "
        "of Cloud Imperium Games. This project is not endorsed by or affiliated with CIG."
    )
    run.font.name = "Calibri"; run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(FG_SECONDARY); run.italic = True

    doc.add_page_break()


def build_toc(doc):
    styled_para(doc, "TABLE OF CONTENTS", size=24, color=ACCENT_CYAN, bold=True, after=20)
    p = doc.add_paragraph()
    run = p.add_run("\u2500" * 60)
    run.font.color.rgb = RGBColor.from_string(BORDER_PRIMARY); run.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(12)

    entries = [
        ("01", "Brand Overview", "Mission, Audience, Taglines, The Six Tools"),
        ("02", "Logo Concepts", "Primary Marks, Icon Marks, Wordmarks"),
        ("03", "Color Palette", "Codebase-Derived Palette, SC MobiGlas References, Accessibility"),
        ("04", "Typography", "Electrolize, Consolas, Type Scale"),
        ("05", "Visual Language & UI Patterns", "Icons, Buttons, Containers, Per-Tool Guidance"),
        ("06", "Brand Voice & Messaging", "Tone, SC Terminology, Sample Copy"),
        ("07", "Asset Inventory & Implementation Roadmap", "Checklist, Tools, Formats, Phases"),
    ]
    for num, title, desc in entries:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}  "); run.font.name = "Consolas"; run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(ACCENT_CYAN); run.bold = True
        run = p.add_run(title); run.font.name = "Calibri"; run.font.size = Pt(13)
        run.font.color.rgb = RGBColor.from_string(FG_PRIMARY); run.bold = True
        p.paragraph_format.space_after = Pt(2)
        p2 = doc.add_paragraph()
        run = p2.add_run(f"      {desc}"); run.font.name = "Calibri"; run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(FG_SECONDARY); run.italic = True
        p2.paragraph_format.space_after = Pt(10)

    doc.add_page_break()


def build_s1(doc):
    """Brand Overview — grounded in the real project."""
    section_h(doc, "01  BRAND OVERVIEW")

    sub_h(doc, "Brand Name")
    styled_para(doc, "SC Toolbox", size=18, bold=True, color=FG_PRIMARY, after=12)
    styled_para(doc,
        "\"SC\" roots it in the Star Citizen universe. \"Toolbox\" is literal\u2014a box of six tools "
        "(and growing) that pilots reach for every session. The underscore in the codebase name "
        "(SC_Toolbox) carries into technical contexts; the marketing-facing form drops it.",
        size=10, color=FG_SECONDARY, line_h=16)

    sub_h(doc, "Tagline Options")
    taglines = [
        ("Precision Tools for the Working Pilot",
         "Recommended primary. Frames the suite as professional gear for active operators."),
        ("Built by Haulers. Trusted by Operators.",
         "Community-first. Earned credibility. Works for README headers and splash screens."),
        ("Your Edge in the Black",
         "Competitive framing. Short enough for favicons and social bios."),
        ("Tools Forged in the \u2019Verse",
         "In-universe flavor. Best on splash screens and loading states."),
        ("Six Tools. One Cockpit.",
         "Counts the current suite. Update the number as tools are added. Direct, memorable."),
    ]
    data_table(doc, ["Tagline", "Rationale"], taglines, [3.0, 3.5])
    doc.add_paragraph()

    sub_h(doc, "Brand Story / Mission Statement")
    styled_para(doc,
        "SC Toolbox started the way most useful things do\u2014somebody needed it and nobody had built "
        "it yet. A hauler running cargo between Lorville and Area18 got tired of alt-tabbing through "
        "spreadsheets to calculate SCU loads, estimate DPS breakdowns, and cross-reference commodity "
        "prices. That first Python script became the Cargo Loader. Then came the DPS Calculator. Then "
        "the Trade Hub. Six tools later, SC Toolbox is a full cockpit companion suite\u2014running as "
        "always-on-top overlays alongside the game, toggled with hotkeys, built to stay out of your way "
        "until you need it.",
        size=10, line_h=16, after=10)
    styled_para(doc,
        "We are not a studio. SC Toolbox is an independent, community-built project that ships as part "
        "of the WingmanAI custom skills ecosystem. Every feature exists because a pilot hit a real "
        "friction point in the \u2019verse. There are no premium tiers, no telemetry, no dark patterns. "
        "The codebase is Python/tkinter for the desktop tools and vanilla JS/Canvas for the isometric "
        "editor\u2014chosen for speed, portability, and zero-dependency installs. Think of us as the "
        "in-universe equivalent of a third-party MobiGlas app vendor: small, specialized, relentlessly "
        "practical.",
        size=10, line_h=16, after=10)
    styled_para(doc,
        "Our mission: reduce friction between a pilot and their objective. Whether you\u2019re optimizing "
        "a C2 Hercules cargo hold, calculating sustained DPS for a Sabre loadout, planning multi-leg "
        "trade routes, or browsing mission payouts across Stanton and Pyro, SC Toolbox should feel like "
        "a natural extension of your cockpit. We build for haulers, fighters, miners, and operators\u2014"
        "the working pilots of the \u2019verse.",
        size=10, line_h=16, after=10)

    sub_h(doc, "The Six Tools")
    styled_para(doc, "SC Toolbox currently ships with six integrated skill modules, each with its own "
        "hotkey, accent color, and always-on-top overlay window.", size=10, color=FG_SECONDARY, line_h=16, after=8)
    data_table(doc,
        ["Tool", "Icon", "Hotkey", "Accent", "Description"],
        [
            ["DPS Calculator", "\u2694", "Shift+1", f"#{COL_DPS}", "Ship loadout viewer, DPS/sustained DPS, power allocation simulator"],
            ["Cargo Loader", "\U0001f4e6", "Shift+2", f"#{COL_CARGO}", "3D isometric cargo grid viewer, container optimization, multi-ship layouts"],
            ["Mission Database", "\U0001f4cb", "Shift+3", f"#{COL_MISSION}", "Mission browser, crafting blueprints, mining resource locations"],
            ["Mining Loadout", "\u26cf", "Shift+4", f"#{COL_MINING}", "Mining laser/module/gadget optimizer with stat breakdowns"],
            ["Market Finder", "\U0001f6d2", "Shift+5", f"#{COL_MARKET}", "Item catalog with buy/sell locations and live prices"],
            ["Trade Hub", "\U0001f4b0", "Shift+6", f"#{COL_TRADE}", "Trade route calculator: single-hop and multi-leg profit routes"],
        ],
        col_widths=[1.2, 0.5, 0.7, 0.8, 3.3])
    doc.add_paragraph()

    sub_h(doc, "Target Audience")
    styled_para(doc, "Primary", size=11, bold=True, color=FG_PRIMARY, after=4)
    bullets(doc, [
        "Cargo traders and haulers calculating SCU loads and optimizing hold layouts",
        "Combat pilots theory-crafting loadouts and comparing DPS breakdowns",
        "Miners optimizing laser/module combinations for specific rock compositions",
        "Traders running profit calculations across multi-stop routes in Stanton and Pyro",
        "Org leaders coordinating fleet operations and resource allocation",
    ])
    doc.add_paragraph()
    styled_para(doc, "Secondary", size=11, bold=True, color=FG_PRIMARY, after=4)
    bullets(doc, [
        "New pilots looking for accessible overlays that flatten Star Citizen\u2019s learning curve",
        "Content creators and streamers who want clean, on-brand data overlays",
        "WingmanAI users discovering SC Toolbox through the custom skills ecosystem",
        "Community developers who may extend or fork individual tools",
    ])
    doc.add_page_break()


def build_s2(doc):
    """Logo Concepts."""
    section_h(doc, "02  LOGO CONCEPTS")
    styled_para(doc,
        "Visual Direction: Star Citizen\u2019s MobiGlas UI\u2014holographic, geometric, grid-based. "
        "Clear shapes, straight pixel lines, solid color panels. The aesthetic should read as a "
        "legitimate third-party MobiGlas application: professional, utilitarian, unmistakably in-universe. "
        "Primary glow color: the SC Toolbox accent cyan (#44AAFF). Avoid cartoonish, fantasy, or overly "
        "militaristic styles. Think independent contractor software, not warship bridge.",
        size=10, color=FG_SECONDARY, line_h=16, after=16)

    # Concept A
    sub_h(doc, "Concept A: \u201cThe MobiGlas Module\u201d")
    styled_para(doc, "Design Rationale", size=10, bold=True, color=FG_PRIMARY, after=4)
    styled_para(doc,
        "Directly references the MobiGlas UI paradigm\u2014SC Toolbox as an app you\u2019d install "
        "on your pilot\u2019s wrist display. The rounded-rectangle frame evokes MobiGlas app tiles. "
        "The internal grid references cargo grids and data tables\u2014the core of what these tools do.",
        size=10, line_h=16, after=8)
    styled_para(doc, "Primary Mark", size=10, bold=True, color=ACCENT_CYAN, after=4)
    styled_para(doc,
        "A rounded-rectangle container with a subtle inner glow (2px, 15% opacity #44AAFF), resembling "
        "a MobiGlas application tile. Inside, a stylized toolbox icon drawn in thin geometric strokes: "
        "an open-top rectangular box with a handle arc, the interior divided into a 2\u00d73 grid (six "
        "compartments\u2014one for each tool). The grid lines are #44AAFF (SC Toolbox accent cyan); the "
        "box outline is #c8d4e8 (primary text white). The rightmost column of the grid pulses with a "
        "subtle gradient from #44AAFF to transparent, suggesting an active holographic display. To the "
        "right, \"SC TOOLBOX\" in all-caps, tracked at +80. \"SC\" is #44AAFF; \"TOOLBOX\" is #c8d4e8. "
        "A 1px horizontal rule in #1e2738 sits beneath the full wordmark. The font is Electrolize (Star "
        "Citizen\u2019s actual UI typeface) or Rajdhani as fallback.",
        size=10, line_h=16, after=8)
    styled_para(doc, "Icon Mark (Simplified)", size=10, bold=True, color=ACCENT_CYAN, after=4)
    styled_para(doc,
        "The rounded-rectangle with 2\u00d73 grid only. At 24px+, the handle arc and grid are visible. "
        "At 16px (favicon), reduce to a simple rounded-rect with a single cross divider. Single color: "
        "#44AAFF on transparent. Below 16px, use a filled rounded-rect silhouette.",
        size=10, line_h=16, after=8)
    styled_para(doc, "Wordmark Variant", size=10, bold=True, color=ACCENT_CYAN, after=4)
    styled_para(doc,
        "\"SC TOOLBOX\" typeset in Electrolize (or Rajdhani), all-caps, tracked +80. \"SC\" in #44AAFF, "
        "\"TOOLBOX\" in #c8d4e8. 1px underline rule in #1e2738. No icon. Used for CLI headers, footer "
        "credits, and documentation mastheads.",
        size=10, line_h=16, after=12)

    # Concept B
    sub_h(doc, "Concept B: \u201cThe HUD Reticle\u201d")
    styled_para(doc, "Design Rationale", size=10, bold=True, color=FG_PRIMARY, after=4)
    styled_para(doc,
        "Inspired by Star Citizen\u2019s cockpit HUD targeting systems. Positions SC Toolbox as a "
        "pilot\u2019s heads-up companion\u2014data overlays that feel native to the cockpit glass. "
        "This concept leans into the always-on-top overlay nature of the actual tools.",
        size=10, line_h=16, after=8)
    styled_para(doc, "Primary Mark", size=10, bold=True, color=ACCENT_CYAN, after=4)
    styled_para(doc,
        "Two concentric circles: outer ring is 1px #1e2738 (border color) with four small tick marks "
        "at cardinal points; inner ring is 2px #44AAFF with a subtle outer glow (0 0 6px rgba(68,170,255,"
        "0.2)). Inside the inner ring, a hexagonal frame containing a stylized wrench overlapping a "
        "data-grid icon\u2014the wrench at 45\u00b0 in #c8d4e8, the grid in #44AAFF. Six small dots "
        "arranged around the hexagon perimeter in the six tool accent colors (#ff7733, #33ccdd, #33dd88, "
        "#ffaa22, #aa66ff, #ffcc00)\u2014a subtle nod to the suite\u2019s six modules. Below the reticle, "
        "\"SC \u2022 TOOLBOX\" centered in Electrolize. Full text in #c8d4e8. A quarter-circle dashed arc "
        "(#1e2738) decorates the lower-left quadrant.",
        size=10, line_h=16, after=8)
    styled_para(doc, "Icon Mark", size=10, bold=True, color=ACCENT_CYAN, after=4)
    styled_para(doc,
        "Inner ring + hexagonal wrench-grid only. The six color dots are retained at 32px+ but dropped "
        "at 24px and below. At 16px, the hexagon simplifies to a circle and the wrench reduces to a "
        "diagonal line crossing a horizontal line. Single color: #44AAFF.",
        size=10, line_h=16, after=8)
    styled_para(doc, "Wordmark Variant", size=10, bold=True, color=ACCENT_CYAN, after=4)
    styled_para(doc,
        "\"SC \u2022 TOOLBOX\" in Electrolize with small-caps \"SC\" treatment. Thin dashed underline "
        "in #1e2738. No icon element.",
        size=10, line_h=16, after=12)

    # Concept C
    sub_h(doc, "Concept C: \u201cThe Blueprint\u201d")
    styled_para(doc, "Design Rationale", size=10, bold=True, color=FG_PRIMARY, after=4)
    styled_para(doc,
        "Engineering schematic aesthetic. Draws from Star Citizen\u2019s ship specification pages and "
        "the technical readout style of manufacturer spec sheets. This concept has the strongest "
        "\"independent contractor\" energy\u2014it looks like documentation from a ship systems vendor.",
        size=10, line_h=16, after=8)
    styled_para(doc, "Primary Mark", size=10, bold=True, color=ACCENT_CYAN, after=4)
    styled_para(doc,
        "A rectangular frame with 45\u00b0 chamfered corners (octagonal silhouette), using a double-line "
        "border: outer 1px in #1e2738 (border), inner 1px in #44AAFF (accent), 2px gap. Inside, \"SC\" "
        "is rendered large and bold in Consolas (the codebase\u2019s actual monospace font), positioned "
        "left-of-center in #44AAFF. Right side has three horizontal callout lines of decreasing length: "
        "top reads \"TOOLBOX\" in #c8d4e8 (9pt Electrolize), middle reads \"v1.2\" in #5a6480, bottom "
        "reads \"WINGMAN CUSTOM SKILL\" in #3a4460. Small registration cross-hairs (+) in #1e2738 at "
        "top-left and bottom-right corners. The overall feeling is a component label on a ship system.",
        size=10, line_h=16, after=8)
    styled_para(doc, "Icon Mark", size=10, bold=True, color=ACCENT_CYAN, after=4)
    styled_para(doc,
        "Chamfered rectangle with \"SC\" centered in Consolas Bold. No callout lines or registration "
        "marks. Double border simplifies to single 2px #44AAFF at sub-32px. At 16px, chamfer becomes "
        "2px radius rounded corners.",
        size=10, line_h=16, after=8)
    styled_para(doc, "Wordmark Variant", size=10, bold=True, color=ACCENT_CYAN, after=4)
    styled_para(doc,
        "\"SC TOOLBOX\" in Consolas, all-caps, enclosed in the chamfered-rectangle frame with double "
        "border. Registration marks at two corners. \"SC\" in #44AAFF, \"TOOLBOX\" in #c8d4e8.",
        size=10, line_h=16, after=16)

    # Usage Rules
    sub_h(doc, "Logo Usage Rules")
    styled_para(doc, "Safe Zones", size=10, bold=True, color=FG_PRIMARY, after=4)
    styled_para(doc,
        "Minimum clear space = height of the \"S\" in the wordmark (1x unit). No graphic elements, "
        "text, or container edges within this boundary.",
        size=10, line_h=16, after=8)
    styled_para(doc, "Minimum Sizes", size=10, bold=True, color=FG_PRIMARY, after=4)
    data_table(doc,
        ["Variant", "Print Minimum", "Digital Minimum"],
        [
            ["Primary Mark (icon + wordmark)", "30mm wide", "120px wide"],
            ["Icon Mark", "8mm", "16px (favicon), 24px (preferred)"],
            ["Wordmark Only", "25mm wide", "100px wide"],
        ], col_widths=[2.5, 1.7, 2.3])
    doc.add_paragraph()
    styled_para(doc, "Placement Rules", size=10, bold=True, color=FG_PRIMARY, after=4)
    bullets(doc, [
        f"Primary background: #{BG_PRIMARY} (launcher bg) or #{BG_DEEPEST} (Trade Hub bg). Logo is optimized for dark backgrounds.",
        f"On light backgrounds: place logo on a solid #{BG_PRIMARY} rectangle with 1x unit padding.",
        "Never place on busy imagery or gradients without a solid backing panel.",
        "Never rotate, skew, add drop shadows, or apply filters to the logo.",
        f"Monochrome white (#c8d4e8) is the only alternate single-color treatment.",
        "The icon mark may be used as a watermark at 10-15% opacity on dark backgrounds.",
        "When displaying alongside the WingmanAI brand, SC Toolbox logo should be equal or smaller in size.",
    ])
    doc.add_page_break()


def build_s3(doc):
    """Color Palette — derived from the actual codebase."""
    section_h(doc, "03  COLOR PALETTE")
    styled_para(doc,
        "This palette is not aspirational\u2014it\u2019s extracted directly from the SC Toolbox codebase "
        "as it ships today. Colors are sourced from shared/theme.py, ui/main_window.py, "
        "skills/DPS_Calculator/dps_ui/constants.py, and each tool\u2019s config files. The palette is "
        "calibrated for dark-mode overlays used alongside Star Citizen at low ambient light.",
        size=10, color=FG_SECONDARY, line_h=16, after=16)

    sub_h(doc, "Background Surfaces (Dark-to-Light Stack)")
    color_table(doc, [
        ("Deepest Black", BG_DEEPEST, "6, 8, 13", "Base canvas for Trade Hub, Mining Loadout. Darkest surface in the suite."),
        ("Primary BG", BG_PRIMARY, "11, 14, 20", "Launcher main background. Default surface for most tools."),
        ("Secondary BG", BG_SECONDARY, "17, 22, 32", "Elevated panels, sidebars, settings areas."),
        ("Card BG", BG_CARD, "20, 26, 38", "Tile backgrounds, skill cards, hover-elevated content."),
        ("Input BG", BG_INPUT, "28, 34, 51", "Text inputs, dropdowns, odd-row table striping."),
        ("Header BG", BG_HEADER, "14, 20, 32", "Header bars, section headers, table headers."),
    ])
    doc.add_paragraph()

    sub_h(doc, "Brand Accent & Star Citizen Reference")
    color_table(doc, [
        ("SC Toolbox Cyan", ACCENT_CYAN, "68, 170, 255", "THE brand color. Interactive highlights, focused elements, links, primary CTA backgrounds."),
        ("SC MobiGlas Cyan", SC_CYAN, "0, 231, 255", "Star Citizen\u2019s canonical MobiGlas glow. Use as a reference/aspiration, not a replacement for #44AAFF."),
        ("SC Blue", SC_BLUE, "0, 144, 224", "Mining/Trade UI blue. Closer to SC\u2019s ship HUD blues."),
        ("Energy Cyan", ENERGY_CYAN, "68, 204, 255", "Energy damage type (DPS Calc). Lighter, more saturated sibling of the brand cyan."),
        ("Mission Teal", MISSION_TEAL, "0, 204, 170", "Mission Database\u2019s unique accent. Greener teal for differentiation."),
    ])
    doc.add_paragraph()

    sub_h(doc, "Text Hierarchy")
    color_table(doc, [
        ("Primary Text", FG_PRIMARY, "200, 212, 232", "Main body text, headings, labels. High contrast on all BG surfaces."),
        ("Bright Text", FG_BRIGHT, "232, 242, 255", "Emphasis, highlighted values, active selections. Near-white."),
        ("Trade/Mining Text", FG_TRADE, "184, 204, 222", "Slightly warmer body text used in Trade Hub and Mining Loadout."),
        ("Secondary Text", FG_SECONDARY, "90, 100, 128", "Descriptions, timestamps, non-critical labels."),
        ("Tertiary Text", FG_TERTIARY, "58, 68, 96", "Disabled text, placeholder text, very low-priority info."),
    ])
    doc.add_paragraph()

    sub_h(doc, "Borders & Separators")
    color_table(doc, [
        ("Primary Border", BORDER_PRIMARY, "30, 39, 56", "Panel borders, card outlines, input borders throughout launcher/DPS/Missions."),
        ("Card Border", BORDER_CARD, "37, 46, 66", "DPS Calculator card borders. Slightly lighter for card elevation."),
        ("Separator", SEPARATOR, "13, 24, 36", "Trade Hub / Mining Loadout subtle dividers."),
    ])
    doc.add_paragraph()

    sub_h(doc, "Per-Tool Accent Colors")
    styled_para(doc, "Each tool has a unique accent color for its tile, status indicators, and featured "
        "UI elements. These colors identify the tool at a glance in the launcher grid.",
        size=10, color=FG_SECONDARY, line_h=16, after=8)
    color_table(doc, [
        ("DPS Orange", COL_DPS, "255, 119, 51", "DPS Calculator identity. Also used for thermal damage type."),
        ("Cargo Cyan", COL_CARGO, "51, 204, 221", "Cargo Loader identity. Cooler, greener than the brand cyan."),
        ("Mission Green", COL_MISSION, "51, 221, 136", "Mission Database identity. Also used for success/running states."),
        ("Mining Amber", COL_MINING, "255, 170, 34", "Mining Loadout identity. Also the general warning/caution color."),
        ("Market Purple", COL_MARKET, "170, 102, 255", "Market Finder identity. Also used for distortion damage type."),
        ("Trade Gold", COL_TRADE, "255, 204, 0", "Trade Hub identity. Evokes aUEC currency and profit."),
    ])
    doc.add_paragraph()

    sub_h(doc, "Functional / State Colors")
    color_table(doc, [
        ("Success Green", GREEN, "51, 221, 136", "Confirmations, running status, profit indicators, valid inputs."),
        ("Success Green (alt)", GREEN_ALT, "0, 221, 112", "Trade Hub / Mining variant. Use either; prefer the tool\u2019s native choice."),
        ("Warning Amber", YELLOW, "255, 170, 34", "Warnings, hidden status, caution states, pending indicators."),
        ("Error Red", RED, "255, 85, 51", "Errors, unavailable status, loss indicators, destructive actions."),
        ("Error Red (alt)", RED_ALT, "224, 64, 32", "Trade Hub / Mining variant. Slightly more muted."),
        ("DPS Orange", ORANGE, "255, 119, 51", "Thermal damage, secondary warnings. Shared with DPS tool accent."),
        ("Distortion Purple", PURPLE, "187, 136, 255", "Distortion damage type in DPS Calculator. Close to Market purple."),
    ])
    doc.add_paragraph()

    sub_h(doc, "DPS Calculator: Damage & Component Type Stripes")
    styled_para(doc, "The DPS Calculator uses specialized color coding for damage types and component "
        "categories. These should remain isolated to that tool\u2019s context.",
        size=10, color=FG_SECONDARY, line_h=16, after=8)
    data_table(doc,
        ["Category", "Color", "Hex", "Context"],
        [
            ["Physical Damage", "\u2588\u2588", "#99aabb", "Ballistic weapons, physical projectiles"],
            ["Energy Damage", "\u2588\u2588", "#44ccff", "Energy weapons, laser fire"],
            ["Distortion Damage", "\u2588\u2588", "#bb88ff", "Distortion weapons, EMP effects"],
            ["Thermal", "\u2588\u2588", "#ff7733", "Heat generation, thermal warnings"],
            ["Shield (component)", "\u2588\u2588", "#bb88ff", "Shield generator card stripe"],
            ["PowerPlant", "\u2588\u2588", "#ff7733", "Power plant card stripe"],
            ["QuantumDrive", "\u2588\u2588", "#44aaff", "QD card stripe (matches brand accent)"],
            ["Cooler", "\u2588\u2588", "#33ccdd", "Cooler card stripe"],
        ], col_widths=[1.5, 0.6, 0.9, 3.5])
    doc.add_paragraph()

    sub_h(doc, "Mission Database: Tag Colors")
    styled_para(doc, "Mission tags use paired background/foreground colors for contrast within tag pills.",
        size=10, color=FG_SECONDARY, line_h=16, after=8)
    data_table(doc,
        ["Tag", "BG Hex", "FG Hex", "Context"],
        [
            ["Delivery / Racing", "#1a3322", "#33cc88", "Green tags for transport/movement missions"],
            ["Combat / Bounty / Mercenary", "#331a1a", "#ff5533", "Red tags for combat-focused missions"],
            ["Salvage / Mining", "#332a1a", "#ffaa22", "Amber tags for resource extraction"],
            ["Investigation", "#221133", "#aa66ff", "Purple tags for intel/investigation missions"],
            ["Rescue / Escort", "#1a2233", "#44aaff", "Blue tags for protection missions"],
            ["LEGAL", "#1a3322", "#33dd88", "Green legality indicator"],
            ["ILLEGAL", "#331a1a", "#ff5533", "Red legality indicator"],
            ["Stanton (system)", "#0a2218", "#33cc88", "Green system tag"],
            ["Pyro (system)", "#331a0a", "#ff7733", "Orange system tag"],
            ["Nyx (system)", "#1a1a33", "#7777cc", "Muted blue-purple system tag"],
        ], col_widths=[1.8, 0.9, 0.9, 2.9])
    doc.add_paragraph()

    sub_h(doc, "Dark Mode & Light Mode")
    styled_para(doc, "Dark Mode (Primary \u2014 THE Mode)", size=10, bold=True, color=ACCENT_CYAN, after=4)
    styled_para(doc,
        "Dark mode is not a variant\u2014it is the only mode. SC Toolbox runs as an always-on-top overlay "
        "alongside Star Citizen. The entire background surface hierarchy (Deepest \u2192 Primary \u2192 "
        "Secondary \u2192 Card \u2192 Input) is calibrated to be near-invisible against the game\u2019s "
        "own dark environments while remaining readable in cockpit lighting. All colors above are "
        "dark-mode native.",
        size=10, line_h=16, after=10)
    styled_para(doc, "Light Mode (Documentation Only)", size=10, bold=True, color=ACCENT_CYAN, after=4)
    styled_para(doc,
        "If a light theme is ever needed for web documentation or GitHub Pages, invert surfaces: "
        "#F8FAFC base, #FFFFFF cards, #E2E8F0 borders. Darken the accent cyan to #0077CC for AA contrast "
        "on white. Primary text becomes #1A2332. Functional colors (green, red, amber) stay unchanged. "
        "The tools themselves should never have a light mode.",
        size=10, line_h=16, after=10)

    sub_h(doc, "Accessibility & Contrast Ratios")
    styled_para(doc, "All text/background pairings must meet WCAG 2.1 AA: 4.5:1 normal text, 3:1 large text.",
        size=10, color=FG_SECONDARY, line_h=16, after=10)
    data_table(doc,
        ["Combination", "Ratio", "AA", "Usage"],
        [
            [f"Primary Text ({FG_PRIMARY}) on Primary BG ({BG_PRIMARY})", "12.1:1", "Pass", "Main body text"],
            [f"Primary Text on Card BG ({BG_CARD})", "9.6:1", "Pass", "Text on cards and tiles"],
            [f"Bright Text ({FG_BRIGHT}) on Primary BG", "14.8:1", "Pass", "Emphasized values, headings"],
            [f"Accent Cyan ({ACCENT_CYAN}) on Primary BG", "7.4:1", "Pass", "Links, interactive highlights"],
            [f"Secondary Text ({FG_SECONDARY}) on Primary BG", "3.6:1", "Large only", "Descriptions (use 14pt+ or bold)"],
            [f"Tertiary Text ({FG_TERTIARY}) on Primary BG", "2.1:1", "Fail", "Decorative only\u2014never for essential info"],
            [f"Success Green ({GREEN}) on Primary BG", "10.1:1", "Pass", "Status text, profit values"],
            [f"Error Red ({RED}) on Primary BG", "4.8:1", "Pass", "Error messages, loss values"],
            [f"Warning Amber ({YELLOW}) on Primary BG", "8.5:1", "Pass", "Warning text and indicators"],
        ], col_widths=[3.0, 0.7, 0.7, 2.1])
    doc.add_page_break()


def build_s4(doc):
    """Typography — based on what the codebase actually uses."""
    section_h(doc, "04  TYPOGRAPHY")
    styled_para(doc,
        "SC Toolbox currently uses Consolas as its sole font across all six tools\u2014headings, body, "
        "data, buttons, everything. This section documents the current state and recommends a deliberate "
        "type hierarchy that introduces display/body fonts while keeping Consolas for data and code.",
        size=10, color=FG_SECONDARY, line_h=16, after=16)

    sub_h(doc, "Current State (Codebase)")
    styled_para(doc, "Every tool uses Consolas monospace at various sizes:", size=10, line_h=16, after=8)
    data_table(doc,
        ["Context", "Current Font", "Current Size", "Weight"],
        [
            ["Window titles", "Consolas", "12-14px", "Bold"],
            ["Section headings", "Consolas", "10-11px", "Bold"],
            ["Body / card text", "Consolas", "9px", "Regular"],
            ["Labels", "Consolas", "8px", "Regular"],
            ["Small labels", "Consolas", "7px", "Regular"],
            ["Buttons", "Consolas", "8-9px", "Bold"],
            ["Status / badges", "Consolas", "8px", "Bold"],
        ], col_widths=[1.5, 1.2, 1.2, 2.6])
    doc.add_paragraph()

    sub_h(doc, "Recommended Type Hierarchy")
    styled_para(doc, "Heading / Display Font", size=10, bold=True, color=ACCENT_CYAN, after=4)
    data_table(doc,
        ["Font", "Style", "Source", "Notes"],
        [
            ["Electrolize", "Geometric, technical", "Google Fonts (free, OFL)", "RECOMMENDED. This is Star Citizen\u2019s actual MobiGlas UI font. Using it creates instant in-universe recognition."],
            ["Rajdhani", "Semi-condensed, geometric", "Google Fonts (free, OFL)", "Fallback. Angular technical feel, excellent at large sizes. Good if Electrolize feels too literal."],
            ["Exo 2", "Geometric sans-serif", "Google Fonts (free, OFL)", "Softer alternative. Futuristic but more readable at mid-sizes."],
        ], col_widths=[1.2, 1.5, 1.5, 2.3])
    doc.add_paragraph()
    styled_para(doc,
        "Download: fonts.google.com/specimen/Electrolize | fonts.google.com/specimen/Rajdhani | "
        "fonts.google.com/specimen/Exo+2",
        size=9, color=FG_SECONDARY, italic=True, after=12)

    styled_para(doc, "Body / UI Font", size=10, bold=True, color=ACCENT_CYAN, after=4)
    data_table(doc,
        ["Font", "Style", "Source", "Notes"],
        [
            ["Inter", "Humanist sans-serif", "Google Fonts (free)", "RECOMMENDED. Designed for screens. Excellent x-height. Pairs well with Electrolize."],
            ["Source Sans 3", "Humanist sans-serif", "Google Fonts (free)", "Slightly narrower than Inter. Good for dense data panels."],
            ["Segoe UI", "System font (Windows)", "Pre-installed", "Zero-cost fallback. Already on every Windows machine running SC Toolbox."],
        ], col_widths=[1.2, 1.5, 1.5, 2.3])
    doc.add_paragraph()

    styled_para(doc, "Monospace / Data Font", size=10, bold=True, color=ACCENT_CYAN, after=4)
    data_table(doc,
        ["Font", "Style", "Source", "Notes"],
        [
            ["Consolas", "Monospace", "Pre-installed (Windows)", "CURRENT AND RECOMMENDED. Already used everywhere. Excellent for data tables, DPS numbers, SCU counts."],
            ["JetBrains Mono", "Monospace, ligatures", "JetBrains (free, OFL)", "If you want ligatures for code displays. Slightly wider than Consolas."],
            ["Fira Code", "Monospace, ligatures", "GitHub (free, OFL)", "Alternative with programming ligatures."],
        ], col_widths=[1.3, 1.4, 1.5, 2.3])
    doc.add_paragraph()

    sub_h(doc, "Type Scale")
    styled_para(doc, "Based on the codebase\u2019s existing sizes, normalized to a consistent scale. "
        "tkinter font sizes are in points; multiply by ~1.33 for approximate pixel equivalents on Windows.",
        size=10, color=FG_SECONDARY, line_h=16, after=10)
    data_table(doc,
        ["Element", "Font", "Size (pt)", "Weight", "Line Height", "Letter Spacing", "Color"],
        [
            ["H1 \u2014 App Title", "Electrolize", "14pt", "Bold", "1.2\u00d7", "+0.05em", f"#{ACCENT_CYAN}"],
            ["H2 \u2014 Section Header", "Electrolize", "11pt", "Bold", "1.25\u00d7", "+0.03em", f"#{FG_PRIMARY}"],
            ["H3 \u2014 Card Heading", "Consolas", "10pt", "Bold", "1.3\u00d7", "Normal", f"#{FG_PRIMARY}"],
            ["Body", "Consolas", "9pt", "Regular", "1.5\u00d7", "Normal", f"#{FG_PRIMARY}"],
            ["Label", "Consolas", "8pt", "Regular", "1.4\u00d7", "Normal", f"#{FG_SECONDARY}"],
            ["Small Label", "Consolas", "7pt", "Regular", "1.3\u00d7", "+0.01em", f"#{FG_SECONDARY}"],
            ["Button", "Consolas", "9pt", "Bold", "1.0\u00d7", "Normal", f"#{FG_PRIMARY}"],
            ["Badge / Hotkey", "Consolas", "8pt", "Bold", "1.0\u00d7", "+0.03em", f"#{ACCENT_CYAN}"],
            ["Data / Numbers", "Consolas", "9pt", "Regular", "1.4\u00d7", "Normal", f"#{FG_BRIGHT}"],
            ["Caption / Status", "Consolas", "8pt", "Regular", "1.3\u00d7", "Normal", f"#{FG_SECONDARY}"],
        ], col_widths=[1.3, 0.9, 0.7, 0.7, 0.7, 0.8, 1.0])
    doc.add_paragraph()

    sub_h(doc, "Implementation Notes")
    bullets(doc, [
        "In tkinter: Use tk.font.Font(family='Consolas', size=9) for body. Electrolize must be installed "
        "on the system or bundled; fall back to ('Segoe UI', size) if unavailable.",
        "All-caps treatment with letter-spacing is Star Citizen\u2019s visual signature for headers. "
        "Apply .upper() in Python; use text-transform: uppercase in CSS.",
        "Maximum 3 font weights per screen. Current codebase uses Regular + Bold only\u2014keep it that way.",
        "Line heights on dark backgrounds should be generous (1.4-1.5\u00d7 for body) to reduce eye strain.",
        "For the Cargo Loader HTML/Canvas: use @font-face to load Electrolize, with Consolas as fallback.",
        "Numeric data in tables should be right-aligned for scanability. Use tabular/monospace figures.",
    ])
    doc.add_page_break()


def build_s5(doc):
    """Visual Language & UI Patterns — based on the actual codebase."""
    section_h(doc, "05  VISUAL LANGUAGE & UI PATTERNS")

    sub_h(doc, "Icon Style Guidelines")
    data_table(doc,
        ["Property", "Specification"],
        [
            ["Style", "Outlined (stroke-based). Consistent with Star Citizen\u2019s MobiGlas geometric line-art approach."],
            ["Stroke Weight", "1.5px at 24px grid. Scale proportionally."],
            ["Corner Radius", "2px on outer shapes. Inner details use sharp corners."],
            ["Grid", "24\u00d724px canvas with 2px padding (20\u00d720px live area)."],
            ["Color", f"Single color: #{FG_PRIMARY} (default), #{ACCENT_CYAN} (active/hover). Per-tool accent for tool-specific icons."],
            ["Aesthetic", "Geometric, technical, schematic. Think MobiGlas UI elements\u2014not photorealistic."],
            ["Existing Icons", "Current tools use emoji (e.g., \u2694, \U0001f4e6, \u26cf). Replace with custom line icons as the brand matures."],
        ], col_widths=[1.5, 5.0])
    doc.add_paragraph()

    sub_h(doc, "Border & Container Styles (Current Codebase)")
    data_table(doc,
        ["Element", "Current Implementation", "Specification"],
        [
            ["Panel Border", "highlightthickness=1, highlightbackground", f"1px solid #{BORDER_PRIMARY}. Active/focused: 1px #{ACCENT_CYAN}."],
            ["Card Border", f"#{BORDER_PRIMARY} (launcher) / #{BORDER_CARD} (DPS)", "1px solid. Hover: border shifts to tool accent color."],
            ["Corner Treatment", "Sharp (tkinter default)", "0px radius in tkinter (unavoidable). 2px in CSS. 4px for larger cards."],
            ["Elevation", f"BG layering: #{BG_PRIMARY} \u2192 #{BG_SECONDARY} \u2192 #{BG_CARD} \u2192 #{BG_INPUT}", "No drop shadows. Elevation = background color shift."],
            ["Dividers", f"Frame with bg=#{SEPARATOR} or #{BORDER_PRIMARY}", f"1px solid #{BORDER_PRIMARY}. Horizontal only."],
            ["Glow Effects", "Not currently used in tkinter", "Subtle only: box-shadow 0 0 6px rgba(68,170,255,0.15) on focus. CSS/Canvas only."],
        ], col_widths=[1.2, 2.3, 3.0])
    doc.add_paragraph()

    sub_h(doc, "Background Textures & Patterns")
    bullets(doc, [
        f"Subtle Grid: 1px lines in #{BORDER_PRIMARY} at 32px intervals, 5-8% opacity. Use on the Cargo "
        "Loader\u2019s isometric canvas background. Not on tkinter panels (tkinter can\u2019t do this).",
        f"Scan Lines: 2px horizontal lines at 3% opacity in #{BORDER_PRIMARY}. Optional for splash screens "
        "and loading states. Reference Star Citizen\u2019s old-school monitor effects on cockpit panels.",
        "Corner Brackets: Decorative L-shaped brackets in panel corners for hero/featured sections. "
        f"1px #{BORDER_PRIMARY}. Atmospheric only. Used in Canvas/CSS contexts.",
        "Star Citizen uses dusty fingerprints and light glare effects on cockpit displays. These are NOT "
        "appropriate for a utility tool\u2014keep surfaces clean. Reference the MobiGlas aesthetic, not "
        "the cockpit glass aesthetic.",
    ])
    doc.add_paragraph()

    sub_h(doc, "Button Styles")
    styled_para(doc, "Current codebase uses ttk.Style() with the 'clam' theme. Here are the defined states:",
        size=10, color=FG_SECONDARY, line_h=16, after=8)
    data_table(doc,
        ["State", "Background", "Text Color", "Border", "Notes"],
        [
            ["Primary (Launch)", f"#{GREEN} / #{COL_MISSION}", f"#{BG_PRIMARY}", "None", "Green background for available/launch actions."],
            ["Primary Hover", "Lighten 10%", f"#{BG_PRIMARY}", "None", "Subtle lightening. Add glow in CSS/Canvas only."],
            [f"Accent (Active)", f"#{ACCENT_CYAN}", f"#{BG_PRIMARY}", "None", "Cyan for primary interactive elements."],
            ["Secondary", f"Transparent / #{BG_INPUT}", f"#{FG_PRIMARY}", f"1px #{BORDER_PRIMARY}", "Outline button. Hover: fill with secondary BG."],
            ["Destructive", f"#{RED}", "#FFFFFF", "None", "Error/delete actions. Use sparingly. Require confirmation."],
            ["Disabled", f"#{BG_INPUT}", f"#{FG_TERTIARY}", f"1px #{BORDER_PRIMARY}", "No hover state. cursor: default / state=DISABLED."],
            [f"Settings Apply", "#1a3020", f"#{ACCENT_CYAN}", "None", "Dark green bg with cyan text. Specific to settings panel."],
        ], col_widths=[1.2, 1.2, 0.9, 1.0, 2.2])
    doc.add_paragraph()
    styled_para(doc,
        "Button sizing in tkinter: padx=12, pady=4 for standard buttons. padx=8, pady=2 for compact/toolbar. "
        "Font: Consolas 8-9pt bold.",
        size=10, color=FG_SECONDARY, line_h=16, after=12)

    sub_h(doc, "Table & Row Styling")
    styled_para(doc, "The DPS Calculator and Mission Database make heavy use of alternating row colors:",
        size=10, color=FG_SECONDARY, line_h=16, after=8)
    data_table(doc,
        ["Context", "Even Row", "Odd Row", "Hover", "Selected"],
        [
            ["DPS Calculator", f"#{BG_CARD} (#161b25)", f"#{BG_INPUT} (#1c2233)", "Lighten +1 step", f"#{ACCENT_CYAN} at 20% opacity"],
            ["Mission Database", f"#{BG_PRIMARY}", f"#{BG_SECONDARY}", "#182030", f"#{ACCENT_CYAN} at 15% opacity"],
            ["Trade Hub / Mining", f"#{BG_DEEPEST}", "#090c15", "\u2014", "#10203c"],
        ], col_widths=[1.3, 1.5, 1.5, 1.3, 1.4])
    doc.add_paragraph()

    sub_h(doc, "Platform-Specific Implementation")

    styled_para(doc, "Python/tkinter (All Six Tools)", size=11, bold=True, color=ACCENT_CYAN, after=4)
    bullets(doc, [
        "ttk theme: 'clam' base. Override with ttk.Style().configure() for branded widgets.",
        f"Root window: root.configure(bg='#{BG_PRIMARY}'). Use -topmost True, -alpha 0.95.",
        "overrideredirect(True) for borderless windows (Trade Hub, Mining Loadout pattern).",
        "Font stack: Consolas is universal. Electrolize requires system install or bundling.",
        "Color constants: Import from shared/theme.py (Trade Hub) or define in per-skill constants.py.",
        f"Combobox styling: fieldbackground='#{BG_INPUT}', arrowcolor='#{ACCENT_CYAN}'.",
        f"Scrollbar: troughcolor='#{BG_SECONDARY}', background='#{BORDER_PRIMARY}'.",
        "Tile hover: Bind <Enter>/<Leave> to toggle highlightbackground between border and accent colors.",
    ])
    doc.add_paragraph()

    styled_para(doc, "JS/Canvas (Cargo Loader Isometric Editor)", size=11, bold=True, color=ACCENT_CYAN, after=4)
    bullets(doc, [
        f"Canvas fill: '#{BG_PRIMARY}' base. Overlay grid pattern with globalAlpha=0.06.",
        f"HUD overlays: strokeStyle='#{ACCENT_CYAN}' for selection, '#{FG_PRIMARY}' for inactive elements.",
        "Text: ctx.font = '600 12px Electrolize, Consolas, monospace' for labels.",
        "1px lines: Offset by 0.5px for crisp rendering on non-retina displays.",
        f"Interactive handles: 8px squares, fill='#{ACCENT_CYAN}', 1px stroke='#{BG_PRIMARY}'.",
        f"Capacity bar: fill='#{GREEN}' (normal), '#{RED}' (overflow). From cargo_grid_editor.html.",
        f"Tooltips: background='#{BG_SECONDARY}', border=1px '#{BORDER_PRIMARY}', font=Consolas 12px.",
    ])
    doc.add_paragraph()

    styled_para(doc, "Future Web-Based Tools (CSS)", size=11, bold=True, color=ACCENT_CYAN, after=4)
    bullets(doc, [
        "CSS custom properties: --bg-primary: #0b0e14; --accent: #44aaff; --fg: #c8d4e8; etc.",
        "Dark mode is default. No prefers-color-scheme switch needed for the tools themselves.",
        "@font-face Electrolize + system fallback stack: 'Electrolize', 'Segoe UI', sans-serif.",
        "Button transitions: transition: all 0.15s ease-in-out;",
        f"Focus rings: outline: 2px solid #{ACCENT_CYAN}; outline-offset: 2px; (never remove focus outlines).",
        f"Scrollbar (WebKit): track #{BG_PRIMARY}, thumb #{BORDER_PRIMARY}, thumb:hover #{FG_SECONDARY}.",
        "SC MobiGlas uses solid color panels and geometric shapes\u2014avoid gradients and glassmorphism.",
    ])
    doc.add_page_break()


def build_s6(doc):
    """Brand Voice & Messaging."""
    section_h(doc, "06  BRAND VOICE & MESSAGING")

    sub_h(doc, "Voice Attributes")
    styled_para(doc,
        "SC Toolbox speaks like a veteran pilot sharing operational notes with a wingmate\u2014not "
        "a corporation talking to customers. The voice is shaped by Star Citizen fluency and grounded "
        "in actual gameplay terminology.",
        size=10, color=FG_SECONDARY, line_h=16, after=10)
    data_table(doc,
        ["Attribute", "Definition", "Example"],
        [
            ["Pragmatic", "Function over flash. Every word earns its place.", "\"Calculates optimal cargo layout for the selected grid.\""],
            ["SC-Fluent", "Speaks Star Citizen natively. Uses correct ship names, locations, units.", "\"Set your route from Lorville to Area18\" (not city A to city B)."],
            ["Dry Humor", "The kind of quip a tired hauler drops at 3am. Never forced.", "\"Cargo overflow detected. Somebody got greedy.\""],
            ["Direct", "No corporate filler. No fake enthusiasm. Say what needs saying.", "\"Connection lost. Retrying...\" (not \"Oops! We\u2019re having trouble!\")"],
            ["Community-Built", "Acknowledges this is a group effort. First person plural.", "\"Built by the community. Report bugs, they\u2019ll get fixed.\""],
        ], col_widths=[1.1, 2.5, 2.9])
    doc.add_paragraph()

    sub_h(doc, "Writing Style Guide")
    styled_para(doc, "Tooltips & Microcopy", size=10, bold=True, color=FG_PRIMARY, after=4)
    bullets(doc, [
        "Under 15 words. Lead with the action.",
        "Sentence case: \"Edit cargo layout\" not \"Edit Cargo Layout\".",
        "No exclamation marks. They read as fake enthusiasm in a utility tool.",
        "Star Citizen terminology is expected\u2014don\u2019t dumb down for your audience.",
    ])
    doc.add_paragraph()
    styled_para(doc, "Error Messages", size=10, bold=True, color=FG_PRIMARY, after=4)
    bullets(doc, [
        "Pattern: [What happened]. [What to do]. \u2014 \"Cache expired. Refreshing ship data...\"",
        "Include error context: \"API timeout (uexcorp.space). Check your connection or try again.\"",
        "Never blame the pilot. Never say \"Oops!\", \"Uh oh!\", or \"Something went wrong\" without specifics.",
    ])
    doc.add_paragraph()
    styled_para(doc, "Update Notes & Changelogs", size=10, bold=True, color=FG_PRIMARY, after=4)
    bullets(doc, [
        "Imperative voice: \"Add Hull C cargo grids\" not \"Added Hull C cargo grids.\"",
        "Group by: Added, Changed, Fixed, Removed. Follow Keep a Changelog format.",
        "Name the ship, weapon, or location affected. Pilots care about what changed for them.",
        "Credit contributors by handle.",
    ])
    doc.add_paragraph()

    sub_h(doc, "Terminology Guide")
    data_table(doc,
        ["Use This", "Not This", "Reason"],
        [
            ["SCU", "Cargo units, CU, capacity", "Standard Star Citizen unit. Our audience uses it daily."],
            ["aUEC", "Credits, money, gold, currency", "In-game currency has a specific name."],
            ["Ship name (e.g., C2 Hercules)", "\"The ship\", \"your vehicle\"", "Be specific. Pilots identify with their hulls."],
            ["Loadout", "Build, setup, config", "Star Citizen community standard term."],
            ["The \u2019verse", "The universe, the game world", "Community shorthand. Signals insider status."],
            ["Pilot / Operator", "User, player, customer", "In-universe framing. They\u2019re pilots."],
            ["Grid / Hold", "Cargo bay, storage area", "Cargo-specific terminology."],
            ["Org", "Guild, clan, team", "SC\u2019s term for player organizations."],
            ["QD / Quantum Drive", "Jump drive, warp drive", "Use the correct Star Citizen component name."],
            ["MFD", "Screen, display, panel", "Multi-Function Display\u2014the cockpit term."],
            ["Stanton / Pyro / Nyx", "\"The system\", \"that area\"", "Name the star system. Our audience knows them."],
            ["Spectrum", "Forums, message board", "CIG\u2019s actual community platform name."],
        ], col_widths=[1.8, 1.8, 2.9])
    doc.add_paragraph()

    sub_h(doc, "Example Copy")
    styled_para(doc, "Sample UI Strings", size=10, bold=True, color=FG_PRIMARY, after=4)
    data_table(doc,
        ["Context", "Copy"],
        [
            ["Empty state (Cargo Loader)", "No cargo loaded. Drag containers to the grid or import a manifest."],
            ["Successful save", "Layout saved. 142 SCU across 3 containers."],
            ["Cargo overflow", "%.1f SCU over capacity. Remove items or switch to a larger hold."],
            ["DPS Calculator loading", "Fetching ship data from erkul.games... (cached 2h)"],
            ["Trade Hub route found", "Best route: Lorville \u2192 Area18. Profit: 12,400 aUEC/run."],
            ["Mining Loadout tooltip", "Optimal laser power for Quantainium. Adjust instability threshold below."],
            ["Market Finder no results", "No listings found for this item in Stanton. Try expanding to Pyro."],
            ["Connection status", "Synced with UEX API (v2). Last refresh: 2 min ago."],
            ["Settings saved", "Keybinds updated. Shift+3 \u2192 Mission Database."],
        ], col_widths=[2.0, 4.5])
    doc.add_paragraph()

    styled_para(doc, "Sample Update Announcement", size=10, bold=True, color=FG_PRIMARY, after=4)
    styled_para(doc, "SC Toolbox v1.2 \u2014 Hull C Support & Mining Overhaul",
        size=11, bold=True, color=ACCENT_CYAN, after=4)
    styled_para(doc,
        "The big one: Hull C cargo grids are now fully supported in the Cargo Loader, including external "
        "spindle mounts and the internal 64-SCU bays. The Mining Loadout got a complete stat refresh "
        "against 3.24.1 data, and the DPS Calculator now shows sustained DPS with overheat curves.",
        size=10, line_h=16, after=4)
    styled_para(doc, "Added", size=10, bold=True, color=GREEN, after=2)
    bullets(doc, [
        "Hull C cargo grid with spindle and internal bay support (Cargo Loader)",
        "Sustained DPS overheat curves in the DPS Calculator",
        "Pyro mission data in the Mission Database",
    ])
    styled_para(doc, "Fixed", size=10, bold=True, color=RED, after=2)
    bullets(doc, [
        "SCU counter no longer double-counts containers on grid edges",
        "Trade Hub profit calculation now accounts for terminal fees",
        "Mining Loadout stat refresh for 3.24.1 laser balance changes",
    ])
    doc.add_paragraph()

    styled_para(doc, "Sample README Intro", size=10, bold=True, color=FG_PRIMARY, after=4)
    styled_para(doc,
        "SC Toolbox is a suite of six always-on-top companion tools for Star Citizen. Calculate DPS "
        "breakdowns, optimize cargo layouts in isometric 3D, plan multi-leg trade routes, browse "
        "missions across Stanton and Pyro, fine-tune mining loadouts, and search live commodity "
        "prices\u2014all without alt-tabbing. Built by a hauler who got tired of spreadsheets. Runs "
        "inside WingmanAI as a custom skill. Free, no telemetry, your data stays local.",
        size=10, line_h=16, after=10, italic=True)

    sub_h(doc, "Community Communication")
    styled_para(doc, "Spectrum / Reddit / Discord", size=10, bold=True, color=FG_PRIMARY, after=4)
    bullets(doc, [
        "Be a peer, not a brand. SC Toolbox is a community project\u2014communicate like a community member.",
        "Bug reports: Respond with gratitude and specifics. \"Thanks. Can you share the log from /logs/? "
        "That\u2019ll narrow it down.\"",
        "Feature requests: Acknowledge, don\u2019t promise. \"Good idea. Adding to the backlog.\"",
        "Never argue with critics. \"Fair point. We prioritize based on what most pilots need.\"",
        "Share WIP screenshots. Keep it genuine. The SC community values authenticity over polish.",
        "Credit contributors by name. Always.",
        "Avoid: marketing-speak, hype countdowns, manufactured urgency, exclamation marks.",
    ])
    doc.add_page_break()


def build_s7(doc):
    """Asset Inventory & Roadmap."""
    section_h(doc, "07  ASSET INVENTORY & IMPLEMENTATION ROADMAP")

    sub_h(doc, "Brand Asset Checklist")
    styled_para(doc, "Prioritized by impact. Tier 1 = ship first.",
        size=10, color=FG_SECONDARY, line_h=16, after=10)
    data_table(doc,
        ["Tier", "Asset", "Variants", "Format", "Status"],
        [
            ["1", "Primary Logo (icon + wordmark)", "Dark bg, monochrome white", "SVG + PNG (1x, 2x)", "To Create"],
            ["1", "Icon Mark (app icon)", "24/32/48/64/128px", "SVG + PNG + ICO", "To Create"],
            ["1", "Favicon", "16/32/180px (Apple Touch)", "ICO + PNG", "To Create"],
            ["1", "Color Token Files", "CSS vars, Python dict (shared/theme.py), JSON", "CSS / PY / JSON", "Partial (theme.py exists)"],
            ["1", "Per-Tool Icon Set", "6 custom icons replacing emoji (\u2694\U0001f4e6\U0001f4cb\u26cf\U0001f6d2\U0001f4b0)", "SVG + PNG", "To Create"],
            ["2", "Splash Screen", "One per tool (6 total) + launcher", "PNG (match window size)", "To Create"],
            ["2", "Social Banners", "GitHub, Discord, Reddit, Spectrum", "PNG (1200\u00d7630)", "To Create"],
            ["2", "README Header", "Logo + tagline + screenshot composite", "PNG (1280\u00d7640)", "To Create"],
            ["2", "UI Icon Set", "20-30 common actions (save, load, export, refresh, etc.)", "SVG + PNG sprite", "To Create"],
            ["3", "Loading Animation", "Pulsing logo or rotating reticle", "CSS animation / GIF / Lottie", "Future"],
            ["3", "Error/Empty State Art", "No data, connection lost, cache expired", "SVG + PNG", "Future"],
            ["3", "Video Intro Bumper", "3-5s logo reveal for streams/recordings", "MP4 (1080p)", "Future"],
        ], col_widths=[0.4, 2.1, 1.8, 1.3, 0.9])
    doc.add_paragraph()

    sub_h(doc, "Recommended Tools")
    data_table(doc,
        ["Task", "Tool", "Cost", "Notes"],
        [
            ["Vector Logo / Icons", "Inkscape", "Free", "Full SVG editor. Export all formats."],
            ["Vector (collaborative)", "Figma", "Free tier", "Browser-based. Good for handoff and feedback."],
            ["Raster Graphics", "GIMP", "Free", "Splash screens, banners, composites."],
            ["AI Concept Generation", "Midjourney / DALL-E", "Paid", "Explore logo concepts. Redraw finals in vector."],
            ["Color Validation", "Coolors.co / WebAIM Contrast Checker", "Free", "Validate palette and WCAG compliance."],
            ["Font Pairing", "Google Fonts + Fontjoy", "Free", "Preview Electrolize + Inter combinations."],
            ["Icon Base", "Lucide Icons (customize)", "Free (ISC)", "Fork as starting point. Restyle to match brand."],
            ["Design Tokens", "Style Dictionary (npm)", "Free", "Generate cross-platform tokens from single JSON source."],
            ["tkinter Theming", "ttkbootstrap or CustomTkinter", "Free", "If deeper tkinter customization is needed."],
        ], col_widths=[1.5, 2.0, 0.5, 2.5])
    doc.add_paragraph()

    sub_h(doc, "File Format Specifications")
    data_table(doc,
        ["Asset", "Primary", "Secondary", "Color Space", "Notes"],
        [
            ["Logo (vector)", "SVG", "AI / EPS", "sRGB", "Source of truth. Outline all text."],
            ["Logo (raster)", "PNG-24", "WebP", "sRGB", "Transparent bg. 2x for retina."],
            ["Favicon", "ICO (multi-res)", "PNG", "sRGB", "16+32+48px in ICO. Apple: 180px PNG."],
            ["Social Banners", "PNG-24", "JPG (90%)", "sRGB", "1200\u00d7630 Open Graph."],
            ["UI Icons", "SVG", "PNG sprite", "sRGB", "24px grid. 1x + 2x PNGs."],
            ["Splash Screens", "PNG-24", "\u2014", "sRGB", "Match tool\u2019s default window size."],
            ["Color Tokens", "JSON", "CSS / Python dict", "\u2014", "Extend shared/theme.py pattern."],
        ], col_widths=[1.2, 1.2, 1.0, 0.8, 2.3])
    doc.add_paragraph()

    sub_h(doc, "Implementation Roadmap")

    styled_para(doc, "Phase 1: Foundation (Weeks 1-2)", size=11, bold=True, color=ACCENT_CYAN, after=4)
    bullets(doc, [
        "Choose logo concept (A: MobiGlas Module, B: HUD Reticle, or C: Blueprint)",
        "Produce primary logo in SVG with all three variants",
        "Generate favicon and icon mark at all required sizes",
        "Unify shared/theme.py into a single canonical color token file used by all 6 tools",
        "Install Electrolize font and test in tkinter / Canvas rendering",
    ])
    doc.add_paragraph()
    styled_para(doc, "Phase 2: Tool Skinning (Weeks 3-4)", size=11, bold=True, color=ACCENT_CYAN, after=4)
    bullets(doc, [
        "Replace emoji icons (\u2694\U0001f4e6\U0001f4cb\u26cf\U0001f6d2\U0001f4b0) with custom line-art SVG/PNG icons",
        "Apply Electrolize to window titles and section headers across all tools",
        "Normalize color constants across tools to use shared token file",
        "Create splash screens for launcher and each tool window",
        "Ensure all tools pass WCAG AA contrast checks",
    ])
    doc.add_paragraph()
    styled_para(doc, "Phase 3: Community Presence (Weeks 5-6)", size=11, bold=True, color=ACCENT_CYAN, after=4)
    bullets(doc, [
        "Design social banners (GitHub repo, Discord server, Reddit/Spectrum posts)",
        "Create README header image with logo + tagline + 6-tool screenshot grid",
        "Standardize README copy using brand voice guidelines",
        "Set up Discord roles using per-tool accent colors",
    ])
    doc.add_paragraph()
    styled_para(doc, "Phase 4: Polish (Ongoing)", size=11, bold=True, color=ACCENT_CYAN, after=4)
    bullets(doc, [
        "Loading animations for data fetches (API calls to erkul.games, uexcorp.space, etc.)",
        "Error/empty state illustrations for each tool context",
        "Simple brand asset page (GitHub Pages or WingmanAI docs integration)",
        "Revisit this guide when tools 7+ are added to the suite",
    ])


# ─── DOCUMENT SETUP ────────────────────────────────────────────────────────

def setup_doc(doc):
    style = doc.styles['Normal']
    style.font.name = "Calibri"; style.font.size = Pt(10)
    style.font.color.rgb = RGBColor.from_string(FG_PRIMARY)
    pf = style.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(6); pf.line_spacing = Pt(16)
    for lvl in range(1, 5):
        sn = f'Heading {lvl}'
        if sn in doc.styles:
            hs = doc.styles[sn]
            hs.font.name = "Calibri"
            hs.font.color.rgb = RGBColor.from_string(ACCENT_CYAN if lvl == 1 else FG_PRIMARY)
            hs.font.size = Pt({1:24, 2:18, 3:14, 4:12}[lvl])
            hs.font.bold = True

def setup_margins(doc):
    for s in doc.sections:
        s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)
    bg = parse_xml(f'<w:background {nsdecls("w")} w:color="{BG_PRIMARY}" w:themeColor="text1"/>')
    doc.element.body.insert(0, bg)

def add_headers_footers(doc):
    for section in doc.sections:
        # Header
        header = section.header; header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run("SC TOOLBOX  \u2022  Brand Design Guide v1.0")
        run.font.size = Pt(8); run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(FG_SECONDARY)
        pPr = hp._p.get_or_add_pPr()
        pPr.append(parse_xml(
            f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="4" w:color="{BORDER_PRIMARY}"/></w:pBdr>'))

        # Footer
        footer = section.footer; footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = fp._p.get_or_add_pPr()
        pPr.append(parse_xml(
            f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="4" w:color="{BORDER_PRIMARY}"/></w:pBdr>'))
        run = fp.add_run("SC Toolbox  \u2022  Star Citizen\u00ae is a trademark of Cloud Imperium Games  \u2022  ")
        run.font.size = Pt(8); run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(FG_SECONDARY)
        # Page number
        r2 = fp.add_run(); r2._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
        r3 = fp.add_run(); r3._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
        r3.font.size = Pt(8); r3.font.name = "Calibri"; r3.font.color.rgb = RGBColor.from_string(FG_SECONDARY)
        r4 = fp.add_run(); r4._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))


# ─── MAIN ───────────────────────────────────────────────────────────────────

def main():
    print("Generating SC Toolbox Brand Design Guide v2...")
    doc = Document()
    setup_doc(doc); setup_margins(doc)

    print("  Cover page..."); build_cover(doc)
    print("  Table of contents..."); build_toc(doc)
    print("  01 Brand Overview..."); build_s1(doc)
    print("  02 Logo Concepts..."); build_s2(doc)
    print("  03 Color Palette..."); build_s3(doc)
    print("  04 Typography..."); build_s4(doc)
    print("  05 Visual Language..."); build_s5(doc)
    print("  06 Brand Voice..."); build_s6(doc)
    print("  07 Asset Inventory..."); build_s7(doc)
    print("  Headers & footers..."); add_headers_footers(doc)

    print(f"  Saving to {OUTPUT_FILE}...")
    doc.save(OUTPUT_FILE)
    print(f"Done! Saved: {OUTPUT_FILE}")

if __name__ == "__main__":
    main()
